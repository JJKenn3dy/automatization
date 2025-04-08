import docx
import pdfplumber
from PyQt6.QtWidgets import QMessageBox
from docx import Document
from docx.shared import Inches
from docx import Document




def pdf_check(self):
    try:
        # Открытие PDF-файла
        with pdfplumber.open("test.pdf") as pdf:
            text = ""
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        # Создание нового документа Word
        document = Document()
        document.add_paragraph(text)
        document.save("output.docx")
        # Получаем текст из созданного документа
        full_text = getText("output.docx")
        # Вставляем каждую строку как отдельную запись в таблицу keying
        update_keys_table(full_text)
    except FileNotFoundError:
        QMessageBox.warning(self, "Ошибка", "PDF с ключами не был найден")



def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)


    return '\n'.join(fullText)
